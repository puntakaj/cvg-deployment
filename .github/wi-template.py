import os
import sys
import glob
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

doc = Document()
tag = sys.argv[1]
repos_deploy = sys.argv[2]
repos_rollback = sys.argv[3]
has_common_deploy = sys.argv[4]

def read_all_files(folder_path):
    all_files = []
    print(f"Reading all files from: {folder_path}")

    try:
        if not os.path.exists(folder_path):
            print(f"Folder not found: {folder_path}")
            return []
        
        items = os.listdir(folder_path)
        
        for item in items:
            item_path = os.path.join(folder_path, item)
            
            if os.path.isfile(item_path):
                all_files.append(item)
                print(f"File name: {item}")
            else:
                print(f"Folder (skipped): {item}")
    
    except Exception as e:
        print(f"Error: {e}")
    
    print(f"Found {len(all_files)} files total")
    return all_files


# create_sql_table
def create_sql_table(doc, title, records):
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = title
    hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[0].merge(hdr_cells[1])

    row1 = table.rows[1]
    row1.cells[0].text = "SQL Script"
    row1.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row1.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row1.cells[0].paragraphs[0].runs[0].font.bold = True
    row1.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row1.cells[1].text = "Remark"
    row1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row1.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row1.cells[1].paragraphs[0].runs[0].font.bold = True
    row1.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    for filename in records:
        row = table.add_row()
        cell = row.cells[0]
        run = cell.paragraphs[0].add_run(filename)
        run.font.size = Pt(10)


def combine_impact_db(combined_list):
    db_impact = clean_to_db((combined_list))
    return list(set(db_impact))


def clean_to_db(file_list):
    db = []
    for filename in file_list:
        if 'CVG' in filename:
            result = filename[filename.find('CVG'):]
        else:
            result = filename

        dot_pos = result.find('.')
        if dot_pos != -1:
            result = result[:dot_pos]
        
        db.append(result)
    
    return db

# create_impact_table
def create_objective_table(doc, files_dba, files_apo, repos_deploy):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    row1 = table.rows[0]
    row1.cells[0].text = "WR"
    row1.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row1.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row1.cells[0].paragraphs[0].runs[0].font.bold = True
    row1.cells[0].paragraphs[0].runs[0].font.size = Pt(12)

    row1.cells[1].text = "Affected Module"
    row1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row1.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row1.cells[1].paragraphs[0].runs[0].font.bold = True
    row1.cells[1].paragraphs[0].runs[0].font.size = Pt(12)

    row1.cells[2].text = "Description"
    row1.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row1.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row1.cells[2].paragraphs[0].runs[0].font.bold = True
    row1.cells[2].paragraphs[0].runs[0].font.size = Pt(12)

    if files_dba:
        row = table.add_row()
        row.cells[1].text = "Database\n- CONVGPROD"
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[2].text = "เพื่อทำการเพิ่ม structure table"
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    
    if files_apo:
        row = table.add_row()
        row.cells[1].text = "Database\n- CONVGPROD"
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[2].text = "เพื่อทำการเพิ่ม data ในส่วนของ Master Config"
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)

    if repos_deploy:
        row = table.add_row()
        row.cells[1].text = "OpenShift\n- cvg-microservice"
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[2].text = "เพื่อทำการเพิ่ม Feature การทำงานให้รองรับตาม Requirement"
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)


# create_impact_table
def create_impact_table(doc,sir_name,files_impact,repos_deploy):
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    row1 = table.rows[0]
    row1.cells[0].text = "SIR Name"
    row1.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row1.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row1.cells[0].paragraphs[0].runs[0].font.bold = True
    row1.cells[0].paragraphs[0].runs[0].font.size = Pt(12)

    row1.cells[1].text = "Impact"
    row1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row1.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row1.cells[1].paragraphs[0].runs[0].font.bold = True
    row1.cells[1].paragraphs[0].runs[0].font.size = Pt(12)

    row2 = table.rows[1]
    row2.cells[0].text = sir_name
    row2.cells[0].paragraphs[0].runs[0].font.bold = True
    row2.cells[0].paragraphs[0].runs[0].font.size = Pt(12)

    if repos_deploy:
        paragraph_repos = row2.cells[1].paragraphs[0]
        title_micro = paragraph_repos.add_run("Microservice\n")
        title_micro.size = Pt(12)
        title_micro.bold = True
        repos = [item.split(":")[0] for item in repos_deploy.split(";") if item]
        for repo_name in repos:
            print(f"Adding repo impact: {repo_name}")
            repo = paragraph_repos.add_run(f"- {repo_name}\n")
            repo.font.size = Pt(10)
        paragraph_repos.add_run("\n")

    if files_impact and len(files_impact) > 0:
        paragraph_db = row2.cells[1].paragraphs[0]
        title_db = paragraph_db.add_run("Database\n")
        title_db.bold = True
        title_db.size = Pt(12)
        title_db_name = paragraph_db.add_run("CONVGPROD\n")
        title_db_name.size = Pt(10)
        for file_name in files_impact:
            print(f"Adding db impact: {file_name}")
            db = paragraph_db.add_run(f"- {file_name}\n")
            db.font.size = Pt(10)


#create_destination_system_table
def create_destination_system_table(repos_deploy):
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    row1 = table.rows[0]
    row1.cells[0].text = "Server Name"
    row1.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row1.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row1.cells[0].paragraphs[0].runs[0].font.bold = True
    row1.cells[0].paragraphs[0].runs[0].font.size = Pt(12)

    row1.cells[1].text = "Database"
    row1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row1.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row1.cells[1].paragraphs[0].runs[0].font.bold = True
    row1.cells[1].paragraphs[0].runs[0].font.size = Pt(12)

    row1.cells[2].text = "Domain"
    row1.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row1.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row1.cells[2].paragraphs[0].runs[0].font.bold = True
    row1.cells[2].paragraphs[0].runs[0].font.size = Pt(12)

    row1.cells[3].text = "PCR / UR"
    row1.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row1.cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row1.cells[3].paragraphs[0].runs[0].font.bold = True
    row1.cells[3].paragraphs[0].runs[0].font.size = Pt(12)

    row2 = table.rows[1]
    row2.cells[0].text = "OpenShift\n- cvg-microservice "
    row2.cells[0].paragraphs[0].runs[0].font.size = Pt(10)

    row2.cells[1].text = "PCONVGLSTN.lisa.org\n- CONVGPROD "
    row2.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    paragraph_domain = row2.cells[2].paragraphs[0]
    domain_run = paragraph_domain.add_run("cvg-prod.intra.ais")
    domain_run.font.size = Pt(10)
    if 'cvg-app-be' in repos_deploy or 'cvg-app-gui' in repos_deploy:
        extra_run = paragraph_domain.add_run("\ncvg-portal.intra.ais")
        extra_run.font.size = Pt(10)


# create_document_table
def create_document_table(doc):
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    
    row1 = table.rows[0]
    row1.cells[0].text = "Document name :"
    row1.cells[0].paragraphs[0].runs[0].font.bold = True
    row1.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    merged_cell = row1.cells[1].merge(row1.cells[3])
    merged_cell.paragraphs[0].add_run("WI_Convergence_YYYY-MM-DD.docx")
    merged_cell.paragraphs[0].runs[0].font.size = Pt(12)
    
    row2 = table.rows[1]
    row2.cells[0].text = "Created by :"
    row2.cells[0].paragraphs[0].runs[0].font.bold = True
    row2.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    row2.cells[1].text = ""
    row2.cells[2].text = "Created Date :"
    row2.cells[2].paragraphs[0].runs[0].font.bold = True
    row2.cells[2].paragraphs[0].runs[0].font.size = Pt(12)
    row2.cells[3].text = "DD/MM/YYYY"
    row2.cells[3].paragraphs[0].runs[0].font.size = Pt(12)
    
    row3 = table.rows[2]
    row3.cells[0].text = "Company :"
    row3.cells[0].paragraphs[0].runs[0].font.bold = True
    row3.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    row3.cells[1].text = "MIMO Tech."
    row3.cells[1].paragraphs[0].runs[0].font.size = Pt(12)
    row3.cells[2].text = "Department :"
    row3.cells[2].paragraphs[0].runs[0].font.size = Pt(12)
    row3.cells[2].paragraphs[0].runs[0].font.bold = True
    row3.cells[3].text = "BAIC"
    row3.cells[3].paragraphs[0].runs[0].font.size = Pt(12)
    
    row4 = table.rows[3]
    row4.cells[0].text = "On Production Date :"
    row4.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    row4.cells[0].paragraphs[0].runs[0].font.bold = True
    row4.cells[1].text = "DD/MM/YYYY"
    row4.cells[1].paragraphs[0].runs[0].font.size = Pt(12)
    row4.cells[2].text = "Telephone :"
    row4.cells[2].paragraphs[0].runs[0].font.size = Pt(12)
    row4.cells[2].paragraphs[0].runs[0].font.bold = True
    row4.cells[3].text = ""

#create micoerserive deploy and rollback table
def create_repo_table(doc, repos, has_common_deploy):
    if not repos:
        print("No repositories found")
        return

    cvg_app_be = ""
    cvg_app_gui = ""
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    row1 = table.rows[0]
    merged_cell = row1.cells[0].merge(row1.cells[2])
    merged_cell.text = "APP"
    merged_cell.paragraphs[0].runs[0].font.bold = True
    merged_cell.paragraphs[0].runs[0].font.size = Pt(12)
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    row2 = table.rows[1]
    row2.cells[0].text = "Pipeline"
    row2.cells[0].paragraphs[0].runs[0].font.bold = True
    row2.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    row2.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row2.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row2.cells[1].text = "Run workflow"
    row2.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row2.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row2.cells[1].paragraphs[0].runs[0].font.size = Pt(12)
    row2.cells[1].paragraphs[0].runs[0].font.bold = True
    row2.cells[2].text = "Tag"
    row2.cells[2].paragraphs[0].runs[0].font.size = Pt(12)
    row2.cells[2].paragraphs[0].runs[0].font.bold = True
    row2.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row2.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    row3 = table.rows[2]
    row3cell0 = row3.cells[0].paragraphs[0].add_run("PRD - Deploy Microservices")
    row3cell0.font.size = Pt(12)
    row3cell0.font.bold = True
    row3cell1 = row3.cells[1].paragraphs[0].add_run("Y")
    row3cell1.font.size = Pt(12)
    row3cell1.font.bold = True
    row3.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row3.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph_repos = row3.cells[2].paragraphs[0]
    apps_and_tags = [item for item in repos.split(";") if item]
    for app_and_tag in apps_and_tags:
        print(f"Adding repo and tag: {app_and_tag}")
        if app_and_tag.startswith("cvg-app-be"):
            cvg_app_be = app_and_tag
            aAt = paragraph_repos.add_run(f"{app_and_tag}\n")
            aAt.font.size = Pt(10)
        elif app_and_tag.startswith("cvg-app-gui"):
            cvg_app_gui = app_and_tag
        else:
            aAt = paragraph_repos.add_run(f"{app_and_tag}\n")
            aAt.font.size = Pt(10)

    frwText = paragraph_repos.add_run("\nFor Run Workflow\n")
    frwText.font.size = Pt(10)
    frwText.bold = True
    if cvg_app_gui:
        filtered_apps = [item for item in apps_and_tags if not item.startswith("cvg-app-gui")]
        repo = ";".join(filtered_apps) + ";"
        print(f"Updated repo without cvg_app_gui: {repo}")
    else:
        repo = repos
    paragraph_repos.add_run(f"{repo}")

    if cvg_app_be:
        row4 = table.add_row()
        row4.cells[0].text = "PRD - CVG BFF Deploy"
        row4.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        row4.cells[0].paragraphs[0].runs[0].font.bold = True
        row4.cells[1].text = "Y"
        row4.cells[1].paragraphs[0].runs[0].font.size = Pt(12)
        row4.cells[1].paragraphs[0].runs[0].font.bold = True
        row4.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row4.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph_app_be = row4.cells[2].paragraphs[0]
        print(f"Adding repo and tag app be: {cvg_app_be}")
        beAt = paragraph_app_be.add_run(f"{cvg_app_be}\n")
        beAt.font.size = Pt(10)
        frwText = paragraph_app_be.add_run("\nFor Run Workflow\n")
        frwText.font.size = Pt(10)
        frwText.bold = True
        paragraph_app_be.add_run(f"{cvg_app_be.split(':')[1]};")

    if cvg_app_gui:
        row5 = table.add_row()
        row5.cells[0].text = "PRD - CVG Frontend Deploy"
        row5.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        row5.cells[0].paragraphs[0].runs[0].font.bold = True
        row5.cells[1].text = "Y"
        row5.cells[1].paragraphs[0].runs[0].font.size = Pt(12)
        row5.cells[1].paragraphs[0].runs[0].font.bold = True
        row5.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row5.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph_app_gui = row5.cells[2].paragraphs[0]
        print(f"Adding repo and tag app gui: {cvg_app_gui}")
        beAt = paragraph_app_gui.add_run(f"{cvg_app_gui}\n")
        beAt.font.size = Pt(10)
        frwText = paragraph_app_gui.add_run("\nFor Run Workflow\n")
        frwText.font.size = Pt(10)
        frwText.bold = True
        paragraph_app_gui.add_run(f"{cvg_app_gui};")
           
    if has_common_deploy == "true":
        row6 = table.add_row()
        row6.cells[0].text = "PRD - CVG BE Common Deployment"
        row6.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        row6.cells[0].paragraphs[0].runs[0].font.bold = True
        row6.cells[1].text = "Y"
        row6.cells[1].paragraphs[0].runs[0].font.size = Pt(12)
        row6.cells[1].paragraphs[0].runs[0].font.bold = True
        row6.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row6.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph_common = row6.cells[2].paragraphs[0]
        frwText = paragraph_common.add_run("For Run Workflow")
        frwText.font.size = Pt(10)
        frwText.bold = True

def main():
    dba_folders = glob.glob(f"./sprint/{tag}/*/DBA")
    apo_folders = glob.glob(f"./sprint/{tag}/*/APO")

    doc.add_heading("Work Instruction Template", 0)
    create_document_table(doc)

    sir_name = (f"{tag}_Enhance_CVG_Microservice_and_Fix_bug")

    if not apo_folders:
        print("No APO folders found.")
        files_apo = []
    else:
        files_apo = []
        for folder_path_apo in apo_folders:
            print(f"Found APO folder: {folder_path_apo}")
            files_apo.extend(read_all_files(folder_path_apo))

    if not dba_folders:
        print("No DBA folders found.")
        files_dba = []
    else:
        files_dba = []
        for folder_path_dba in dba_folders:
            print(f"Found DBA folder: {folder_path_dba}")
            files_dba.extend(read_all_files(folder_path_dba))

    files_impact = combine_impact_db(files_dba + files_apo)

    doc.add_heading('Objective', level=2)
    create_objective_table(doc, files_dba, files_apo, repos_deploy)

    doc.add_heading('Impact', level=2)
    create_impact_table(doc,sir_name,files_impact,repos_deploy)

    doc.add_heading('Destination System', level=2)
    create_destination_system_table(repos_deploy)

    if files_dba:
        doc.add_heading('Database - DBA', level=2)
        create_sql_table(doc, "DBA", files_dba)

    if files_apo:
        doc.add_heading('Database - APO', level=2)
        create_sql_table(doc, "APO", files_apo)
    
    if repos_deploy:
        doc.add_heading('Application - Operation', level=2)
        create_repo_table(doc, repos_deploy, has_common_deploy)
    
    if repos_rollback:
        doc.add_heading('Rollback', level=2)
        create_repo_table(doc, repos_rollback, has_common_deploy)

    wi_filename = f"wi-{tag}.docx"
    doc.save(wi_filename)
    print(f"Document saved as {wi_filename}")
    

if __name__ == "__main__":
    main()
